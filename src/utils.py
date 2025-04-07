from pathlib import Path
import csv
import contextlib
import io
import shutil

from num2words import num2words
from docx2pdf import convert
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell

from logger import logger, _, lang_name


def replace_words_in_paragraph(
    paragraph: Paragraph,
    replacement_words: dict[str, str],
) -> None:
    for run in paragraph.runs:
        logger.debug(_("Words replaced in run: {text}").format(text=run.text))
        for old_word, new_word in replacement_words.items():
            old_word_fmt = f"[{old_word.upper()}]"
            if old_word_fmt in run.text:
                logger.debug(
                    _('> Word "{old_word}" replaced with "{new_word}"').format(
                        old_word=old_word, new_word=new_word
                    )
                )
                run.text = run.text.replace(old_word_fmt, new_word)


def process_table(
    table: Table,
    replacement_words: dict[str, str],
) -> None:
    for row in table.rows:
        for cell in row.cells:
            process_cell(cell, replacement_words)
            for nested_table in cell.tables:
                process_table(nested_table, replacement_words)


def process_paragraphs(
    paragraphs: list[Paragraph],
    replacement_words: dict[str, str],
) -> None:
    for paragraph in paragraphs:
        replace_words_in_paragraph(
            paragraph,
            replacement_words,
        )


def process_cell(
    cell: _Cell,
    replacement_words: dict[str, str],
) -> None:
    for paragraph in cell.paragraphs:
        replace_words_in_paragraph(
            paragraph,
            replacement_words,
        )


def replace_words_in_doc(
    src_doc_path: Path | str,
    dst_doc_path: Path | str,
    replacement_words: dict[str, str],
) -> None:
    shutil.copy(src_doc_path, dst_doc_path)

    doc = Document(str(dst_doc_path))

    logger.debug(_("Words replaced in paragraphs"))
    process_paragraphs(
        doc.paragraphs,
        replacement_words,
    )

    logger.debug(_("Words replaced in tables"))
    for table in doc.tables:
        process_table(table, replacement_words)

    doc.save(str(dst_doc_path))


def convert_docx_template_to_pdf(
    docx_template_path: Path,
    pdf_file_path: Path,
    replacement_words: dict[str, str],
) -> None:
    docx_modified_path = Path(f"{docx_template_path}.mdf.docx")
    logger.debug(_("File {path} created").format(path=docx_modified_path))

    logger.debug(_("Start words replace in document"))
    replace_words_in_doc(
        docx_template_path,
        docx_modified_path,
        replacement_words,
    )
    logger.debug(_("End words replace in document"))

    logger.debug(_("Start document convert"))
    with (
        contextlib.redirect_stdout(io.StringIO()),
        contextlib.redirect_stderr(io.StringIO()),
    ):
        convert(
            input_path=docx_modified_path,
            output_path=pdf_file_path,
            keep_active=True,
        )
    logger.debug(_("End document convert"))

    docx_modified_path.unlink()
    logger.debug(_("File {path} deleted").format(path=docx_modified_path))

    logger.debug(_("Document converted to {path}").format(path=pdf_file_path))


def number_to_words_currency(
    number: int | float,
    currency_pluralize: list[str],
) -> str:
    if len(currency_pluralize) != 3:
        raise ValueError(
            _("Currency pluralize must be a list with 3 elements")
        )
    if number % 10 == 1 and number % 100 != 11:
        currency = currency_pluralize[0]
    elif 2 <= number % 10 <= 4 and (number % 100 < 10 or number % 100 >= 20):
        currency = currency_pluralize[1]
    else:
        currency = currency_pluralize[2]

    words = num2words(number, lang=lang_name)
    return f"{words} {currency}"


def csv_to_dict(
    csv_path: Path,
    wrong_headers: set[str] | None = None,
) -> dict[str, dict[str, str]]:
    with open(csv_path, "r", encoding="utf-8") as csv_file:
        reader = csv.DictReader(csv_file, delimiter=";")
        headers = reader.fieldnames
        rows = list(reader)
    if not headers:
        raise ValueError(_("Headers not found"))
    if wrong_headers:
        finded_wrong_headers = wrong_headers & set(headers)
        if len(finded_wrong_headers) > 0:
            raise ValueError(
                _("Wrong headers: {headers}").format(
                    headers=", ".join(finded_wrong_headers),
                )
            )
    return {row[headers[0]]: row for row in rows}
